from pathlib import Path

from fastapi.testclient import TestClient

from app.services.document_quality import (
    build_application_answer_sheet,
    run_ats_diagnostics,
    template_config,
    validate_claims,
)


def test_template_config(client: TestClient, auth_headers):
    response = client.get("/api/documents/templates", headers=auth_headers)
    assert response.status_code == 200
    body = response.json()
    assert body["template_version"] == "r2.4.0"
    assert body["generation_mode"] == "deterministic_fallback"


def test_baseline_resume(client: TestClient, auth_headers):
    response = client.get("/api/documents/baseline-resume", headers=auth_headers)
    assert response.status_code == 200
    body = response.json()
    assert body["evidence_count"] >= 1
    assert body["baseline_hash"]


def test_packet_includes_document_quality(client: TestClient, auth_headers):
    job = client.post("/api/jobs/ingest/fixture", headers=auth_headers).json()[0]
    app = client.post(f"/api/applications/jobs/{job['id']}/generate", headers=auth_headers).json()
    quality = app["packet_metadata"]["document_quality"]
    assert quality["template_version"] == "r2.4.0"
    assert quality["ats_diagnostics"]["passed"] is True
    assert quality["answer_sheet"]
    assert app["packet_metadata"]["answer_sheet"]


def test_application_quality_endpoint(client: TestClient, auth_headers):
    job = client.post("/api/jobs/ingest/fixture", headers=auth_headers).json()[0]
    app = client.post(f"/api/applications/jobs/{job['id']}/generate", headers=auth_headers).json()
    response = client.get(f"/api/documents/applications/{app['id']}/quality", headers=auth_headers)
    assert response.status_code == 200
    assert response.json()["passed"] is True


def test_validate_application_documents(client: TestClient, auth_headers):
    job = client.post("/api/jobs/ingest/fixture", headers=auth_headers).json()[0]
    app = client.post(f"/api/applications/jobs/{job['id']}/generate", headers=auth_headers).json()
    response = client.post(f"/api/documents/applications/{app['id']}/validate", headers=auth_headers)
    assert response.status_code == 200
    assert response.json()["passed"] is True
    assert "factual_core" in response.json()


def test_ats_diagnostics_unit(tmp_path):
    from docx import Document

    path = tmp_path / "sample.docx"
    doc = Document()
    doc.add_paragraph("Professional Summary")
    doc.add_paragraph("Summary text")
    doc.add_paragraph("Professional Experience")
    doc.add_paragraph("Experience bullet")
    doc.add_paragraph("Role Target")
    doc.save(path)
    result = run_ats_diagnostics(path)
    assert result["passed"] is True


def test_claim_validation():
    evidence = ["Ascensus Sr Principal SDET – Automation Framework Architect."]
    result = validate_claims("Ascensus automation framework architect", evidence)
    assert result["passed"] is True
    assert result["supported_claims"]


def test_answer_sheet_builder():
    sheet = build_application_answer_sheet(
        job_title="Director QE",
        company="Acme",
        profile_name="QE Leadership",
        keyword_mapping={"automation": ["Built automation platform"]},
        evidence=["Built automation platform"],
    )
    assert "Application Answer Sheet" in sheet
    assert "automation" in sheet
