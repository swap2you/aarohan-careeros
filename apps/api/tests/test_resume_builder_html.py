from pathlib import Path

from docx import Document

from app.services.resume_builder import docx_to_submission_html, extract_docx_plaintext


def test_docx_to_submission_html_excludes_role_target(tmp_path: Path):
    path = tmp_path / "resume.docx"
    doc = Document()
    doc.add_paragraph("Swapnil Patil")
    doc.add_paragraph("Professional Summary")
    doc.add_paragraph("Summary text for the candidate.")
    doc.add_paragraph("Professional Experience")
    doc.add_paragraph("Delivered automation platform outcomes.", style="List Bullet")
    doc.add_paragraph("Role Target")
    doc.add_paragraph("Director QE at Acme")
    doc.save(path)

    html = docx_to_submission_html(path)
    assert "Role Target" not in html
    assert "Professional Summary" in html
    text = extract_docx_plaintext(path)
    assert "Swapnil Patil" in text
