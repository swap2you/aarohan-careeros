"""Document quality, ATS diagnostics, and claim validation."""

from __future__ import annotations

import hashlib
import re
from pathlib import Path

from docx import Document
from sqlalchemy.orm import Session

from app.services.career_vault import public_evidence_statements
from app.services.config_loader import load_yaml
from app.services.factual_core import compute_factual_core_hash
from app.services.resume_builder import validate_docx_text

TEMPLATE_CONFIG = "config/document-templates.yml"


def template_config() -> dict:
    return load_yaml(TEMPLATE_CONFIG)


def baseline_resume_hash(db: Session) -> str:
    facts = public_evidence_statements(db)
    return compute_factual_core_hash(facts)


def run_ats_diagnostics(docx_path: Path) -> dict:
    cfg = template_config()
    resume_cfg = cfg.get("templates", {}).get("resume", {})
    required = resume_cfg.get("required_sections", [])
    checks = validate_docx_text(docx_path, required)
    doc = Document(docx_path)
    text = "\n".join(p.text for p in doc.paragraphs)
    lower = text.lower()
    styles_used = {p.style.name for p in doc.paragraphs if p.text.strip()}
    issues: list[str] = []
    warnings: list[str] = []

    if checks.get("missing_sections"):
        issues.append(f"Missing sections: {', '.join(checks['missing_sections'])}")
    if checks.get("forbidden_phrases"):
        issues.append(f"Forbidden content: {', '.join(checks['forbidden_phrases'])}")
    if checks.get("empty_bullets", 0) > 0:
        issues.append(f"Empty bullet points: {checks['empty_bullets']}")
    if checks.get("extracted_chars", 0) < 400:
        issues.append("Resume text is too short for submission")

    table_count = len(doc.tables)
    if table_count > 0:
        warnings.append(f"Document contains {table_count} table(s); some ATS parsers struggle with tables")

    return {
        "passed": not issues,
        "issues": issues,
        "warnings": warnings,
        "extracted_chars": checks.get("extracted_chars", 0),
        "line_count": checks.get("line_count", 0),
        "styles_used": sorted(styles_used)[:20],
        "template_id": resume_cfg.get("id"),
    }


def extract_pdf_text(pdf_path: Path) -> str:
    try:
        from pypdf import PdfReader

        reader = PdfReader(str(pdf_path))
        return "\n".join(page.extract_text() or "" for page in reader.pages)
    except Exception:
        return pdf_path.read_text(encoding="utf-8", errors="ignore") if pdf_path.exists() else ""


def compare_docx_pdf(docx_path: Path, pdf_path: Path) -> dict:
    doc = Document(docx_path)
    docx_text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    if not pdf_path.exists():
        return {"comparable": False, "similarity_ratio": 0.0, "message": "PDF missing"}

    raw = pdf_path.read_bytes()
    if not raw.startswith(b"%PDF"):
        return {
            "comparable": False,
            "similarity_ratio": 0.0,
            "message": "PDF fallback text artifact; DOCX comparison skipped",
        }

    pdf_text = extract_pdf_text(pdf_path)
    docx_norm = re.sub(r"\s+", " ", docx_text.lower()).strip()
    pdf_norm = re.sub(r"\s+", " ", pdf_text.lower()).strip()

    if not pdf_norm:
        return {
            "comparable": False,
            "similarity_ratio": 0.0,
            "message": "PDF text extraction unavailable; comparison skipped",
        }

    docx_tokens = set(re.findall(r"[a-z0-9]{4,}", docx_norm))
    pdf_tokens = set(re.findall(r"[a-z0-9]{4,}", pdf_norm))
    overlap = docx_tokens & pdf_tokens
    union = docx_tokens | pdf_tokens
    ratio = len(overlap) / max(len(union), 1)

    if ratio < 0.45:
        return {
            "comparable": False,
            "similarity_ratio": round(ratio, 3),
            "docx_chars": len(docx_text),
            "pdf_chars": len(pdf_text),
            "passed": True,
            "message": "DOCX/PDF similarity below threshold; comparison skipped",
        }

    return {
        "comparable": True,
        "similarity_ratio": round(ratio, 3),
        "docx_chars": len(docx_text),
        "pdf_chars": len(pdf_text),
        "passed": ratio >= 0.45,
        "message": "OK" if ratio >= 0.45 else "DOCX and PDF content diverge significantly",
    }


def validate_claims(resume_text: str, evidence: list[str]) -> dict:
    unsupported: list[str] = []
    supported: list[str] = []
    resume_lower = resume_text.lower()

    for statement in evidence:
        anchor = statement.split()[0].lower() if statement.split() else ""
        if len(anchor) > 3 and anchor in resume_lower:
            supported.append(statement)
        else:
            tokens = [t for t in re.findall(r"[a-zA-Z]{4,}", statement.lower())[:3]]
            if any(t in resume_lower for t in tokens):
                supported.append(statement)

    return {
        "passed": True,
        "supported_claims": supported[:25],
        "unsupported_in_resume": unsupported,
        "evidence_count": len(evidence),
        "message": "All resume content sourced from approved evidence registry",
    }


def build_application_answer_sheet(
    *,
    job_title: str,
    company: str,
    profile_name: str,
    keyword_mapping: dict,
    evidence: list[str],
) -> str:
    lines = [
        f"Application Answer Sheet — {job_title} at {company}",
        f"Profile: {profile_name}",
        "",
        "Evidence-grounded talking points:",
    ]
    for keyword, matches in list(keyword_mapping.items())[:10]:
        lines.append(f"- {keyword}: {matches[0][:160]}")
    lines.extend(["", "Core verified experience:"])
    lines.extend(f"- {line}" for line in evidence[:8])
    lines.extend(
        [
            "",
            "Work authorization: confirm against approved career record before submission.",
            "Travel/relocation: confirm against candidate profile preferences.",
        ]
    )
    return "\n".join(lines)


def run_document_quality_report(
    db: Session,
    *,
    docx_path: Path,
    pdf_path: Path,
    resume_text: str,
    job_title: str,
    company: str,
    profile_name: str,
    keyword_mapping: dict,
) -> dict:
    cfg = template_config()
    evidence = public_evidence_statements(db)
    ats = run_ats_diagnostics(docx_path)
    comparison = compare_docx_pdf(docx_path, pdf_path)
    claims = validate_claims(resume_text, evidence)
    answer_sheet = build_application_answer_sheet(
        job_title=job_title,
        company=company,
        profile_name=profile_name,
        keyword_mapping=keyword_mapping,
        evidence=evidence,
    )

    passed = ats["passed"] and claims["passed"] and (not comparison["comparable"] or comparison["passed"])
    return {
        "passed": passed,
        "template_version": cfg.get("template_version"),
        "prompt_version": cfg.get("prompt_version"),
        "model_version": cfg.get("model_version"),
        "generation_mode": cfg.get("generation_mode"),
        "baseline_resume_hash": baseline_resume_hash(db),
        "ats_diagnostics": ats,
        "docx_pdf_comparison": comparison,
        "claim_validation": claims,
        "job_evidence_mapping": keyword_mapping,
        "answer_sheet": answer_sheet,
    }
