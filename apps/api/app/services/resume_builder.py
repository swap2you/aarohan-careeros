"""Professional ATS-safe resume and cover letter builders."""

from __future__ import annotations

import re
from datetime import datetime
from pathlib import Path

import yaml
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt

from app.config import settings

FORBIDDEN_DOCX_PHRASES = (
    "internal review notes",
    "role target",
    "evidence_id",
    "evidence id",
    "remove before submission",
)

REQUIRED_SECTIONS = (
    "Professional Summary",
    "Core Competencies",
    "Professional Experience",
    "Education",
)


def _vault_root() -> Path:
    root = Path(settings.career_vault_root)
    return root if root.exists() else Path(__file__).resolve().parents[4] / "career_vault"


def load_resume_profile(profile_id: str) -> dict:
    path = _vault_root() / "resume_profiles" / f"{profile_id}.yml"
    if not path.exists():
        raise ValueError(f"Unknown resume profile: {profile_id}")
    return yaml.safe_load(path.read_text(encoding="utf-8")) or {}


def extract_keywords(text: str | None) -> list[str]:
    if not text:
        return []
    tokens = re.findall(r"[a-zA-Z][a-zA-Z0-9+#./-]{2,}", text.lower())
    stop = {"the", "and", "for", "with", "you", "our", "will", "this", "that", "from", "have", "your"}
    keywords: list[str] = []
    for token in tokens:
        if token in stop or token in keywords:
            continue
        keywords.append(token)
    return keywords[:40]


def map_keywords_to_evidence(keywords: list[str], evidence: list[str]) -> dict:
    mapping: dict[str, list[str]] = {}
    for keyword in keywords:
        matches = [line for line in evidence if keyword in line.lower()]
        if matches:
            mapping[keyword] = matches[:2]
    return mapping


def _add_heading(doc: Document, text: str) -> None:
    para = doc.add_paragraph(text)
    para.runs[0].bold = True
    para.runs[0].font.size = Pt(12)


def _group_evidence_by_employer(evidence: list[str]) -> list[tuple[str, list[str]]]:
    """Group flat evidence lines into employer blocks when prefixed with employer markers."""
    groups: list[tuple[str, list[str]]] = []
    current_employer = "Professional Experience"
    current_lines: list[str] = []
    for line in evidence:
        stripped = line.strip()
        if not stripped:
            continue
        if stripped.startswith("@"):
            if current_lines:
                groups.append((current_employer, current_lines))
            current_employer = stripped.lstrip("@").strip() or "Professional Experience"
            current_lines = []
        else:
            current_lines.append(stripped)
    if current_lines:
        groups.append((current_employer, current_lines))
    if not groups:
        groups.append(("Professional Experience", evidence[:12]))
    return groups


def validate_docx_text(path: Path, expected_sections: list[str] | None = None) -> dict:
    doc = Document(path)
    text = "\n".join(p.text for p in doc.paragraphs)
    lower = text.lower()
    sections = expected_sections or list(REQUIRED_SECTIONS)
    missing = [section for section in sections if section.lower() not in lower]
    forbidden = [phrase for phrase in FORBIDDEN_DOCX_PHRASES if phrase in lower]
    empty_bullets = sum(1 for p in doc.paragraphs if p.style and "List" in p.style.name and not p.text.strip())
    return {
        "extracted_chars": len(text),
        "missing_sections": missing,
        "forbidden_phrases": forbidden,
        "empty_bullets": empty_bullets,
        "line_count": len(text.splitlines()),
    }


def build_ats_docx(
    *,
    output_path: Path,
    contact: dict,
    profile: dict,
    evidence: list[str],
    job_title: str,
    company: str,
    keyword_mapping: dict,
    missing_warnings: list[str],
) -> dict:
    """Build submission-quality resume. Internal warnings are never written to the document."""
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    title = doc.add_paragraph(contact.get("name", "Candidate"))
    title.runs[0].bold = True
    title.runs[0].font.size = Pt(16)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    contact_line = " | ".join(
        filter(
            None,
            [
                contact.get("email", ""),
                contact.get("phone", ""),
                contact.get("location", ""),
                contact.get("linkedin", ""),
            ],
        )
    )
    if contact_line:
        contact_para = doc.add_paragraph(contact_line)
        contact_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_paragraph("")
    _add_heading(doc, "Professional Summary")
    summary = profile.get("summary", "").strip()
    if profile.get("headline"):
        doc.add_paragraph(profile["headline"].strip())
    if summary:
        doc.add_paragraph(summary)

    competencies = profile.get("core_competencies") or list(keyword_mapping.keys())[:12]
    if not competencies:
        competencies = ["Quality Engineering", "Test Automation", "CI/CD", "Leadership", "API Testing"]
    doc.add_paragraph("")
    _add_heading(doc, "Core Competencies")
    doc.add_paragraph(", ".join(str(c) for c in competencies[:18]))

    technical = profile.get("technical_skills") or []
    if technical:
        doc.add_paragraph("")
        _add_heading(doc, "Technical Skills")
        doc.add_paragraph(", ".join(str(s) for s in technical[:24]))

    doc.add_paragraph("")
    _add_heading(doc, "Professional Experience")
    for employer, lines in _group_evidence_by_employer(evidence):
        employer_para = doc.add_paragraph(employer)
        employer_para.runs[0].bold = True
        for line in lines:
            clean = re.sub(r"^\[?\w+-\d+\]?\s*", "", line).strip()
            if clean:
                doc.add_paragraph(clean, style="List Bullet")

    education = contact.get("education") or profile.get("education") or []
    doc.add_paragraph("")
    _add_heading(doc, "Education")
    if education:
        for item in education:
            if isinstance(item, dict):
                doc.add_paragraph(
                    f"{item.get('degree', '')} — {item.get('institution', '')} ({item.get('year', '')})".strip(" — ()"),
                )
            else:
                doc.add_paragraph(str(item))
    else:
        doc.add_paragraph("See verified career record.")

    certs = contact.get("certifications") or []
    if certs:
        doc.add_paragraph("")
        _add_heading(doc, "Certifications")
        for cert in certs:
            doc.add_paragraph(f"{cert.get('name')} ({cert.get('date', '')})", style="List Bullet")

    doc.core_properties.title = f"Resume — {company} — {job_title}"
    doc.core_properties.author = contact.get("name", "Candidate")
    doc.save(output_path)

    checks = validate_docx_text(output_path, list(REQUIRED_SECTIONS))
    checks["page_warning"] = checks["line_count"] > 120
    checks["internal_warnings_count"] = len(missing_warnings)
    return checks


def build_cover_letter_docx(
    *,
    output_path: Path,
    contact: dict,
    profile: dict,
    job_title: str,
    company: str,
) -> None:
    doc = Document()
    doc.add_paragraph(contact.get("name", "Candidate"))
    doc.add_paragraph(datetime.utcnow().strftime("%B %d, %Y"))
    doc.add_paragraph("")
    doc.add_paragraph(f"Hiring Team\n{company}")
    doc.add_paragraph("")
    doc.add_paragraph(f"Dear Hiring Team at {company},")
    doc.add_paragraph(
        f"I am writing to express my interest in the {job_title} role. "
        f"{profile.get('summary', '').strip()} "
        "I would welcome a conversation about how I can contribute to your team's quality and delivery outcomes."
    )
    doc.add_paragraph("")
    doc.add_paragraph("Regards,")
    doc.add_paragraph(contact.get("name", "Candidate"))
    doc.save(output_path)


def build_interview_brief_docx(
    *,
    output_path: Path,
    contact: dict,
    job_title: str,
    company: str,
    talking_points: list[str],
    gap_notes: list[str] | None = None,
) -> None:
    """Interview preparation brief — no internal IDs or review notes."""
    doc = Document()
    doc.add_paragraph(f"Interview Preparation Brief — {job_title} at {company}")
    doc.add_paragraph(contact.get("name", "Candidate"))
    doc.add_paragraph("")
    _add_heading(doc, "Role Summary")
    doc.add_paragraph(f"Target role: {job_title} at {company}")
    doc.add_paragraph("")
    _add_heading(doc, "Evidence-Grounded Talking Points")
    for point in talking_points[:12]:
        if point.strip():
            doc.add_paragraph(point.strip(), style="List Bullet")
    if gap_notes:
        doc.add_paragraph("")
        _add_heading(doc, "Areas to Clarify")
        for note in gap_notes[:6]:
            if note.strip():
                doc.add_paragraph(note.strip(), style="List Bullet")
    doc.core_properties.title = f"Interview Brief — {company}"
    doc.save(output_path)
